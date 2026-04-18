import fitz  # PyMuPDF
import os

def extract_text_from_pdf(pdf_path):
    """
    Belirtilen PDF dosyasından metin içeriğini çıkarır.
    Hata Yönetimi: Dosya bulunamaz, bozuk PDF, boş sayfa durumları ele alınır.
    """
    if not pdf_path or not os.path.exists(pdf_path):
        print(f"⚠️ Uyarı: '{pdf_path}' dosyası bulunamadı.")
        return ""
    
    try:
        doc = fitz.open(pdf_path)
        
        if doc.page_count == 0:
            print(f"⚠️ Uyarı: '{pdf_path}' dosyası boş (0 sayfa).")
            doc.close()
            return ""
        
        text = ""
        links = []
        
        for page in doc:
            try:
                page_text = page.get_text()
                if page_text:
                    text += page_text
            except Exception as page_err:
                print(f"⚠️ Sayfa okuma hatası (sayfa {page.number}): {page_err}")
                continue
            
            try:
                for link in page.get_links():
                    if 'uri' in link and link['uri']:
                        links.append(link['uri'])
            except Exception:
                pass
        
        doc.close()
        
        # Eğer metin çok kısaysa görsel PDF olabilir → OCR dene
        if len(text.strip()) < 50:
            print(f"📸 '{pdf_path}' görsel PDF olabilir, OCR deneniyor...")
            ocr_text = extract_text_with_ocr(pdf_path)
            if ocr_text and len(ocr_text.strip()) > len(text.strip()):
                text = ocr_text
        
        if links:
            unique_links = list(dict.fromkeys(links))
            text += "\n\n[EK_LINKLER]: " + " ".join(unique_links)
        
        if not text.strip():
            print(f"⚠️ Uyarı: '{pdf_path}' dosyasından metin çıkarılamadı.")
        
        return text
        
    except Exception as e:
        print(f"❌ PDF Hata: {pdf_path} → {e}")
        return ""


def extract_text_from_docx(docx_path):
    """
    DOCX (Word) dosyasından metin çıkarır.
    """
    if not docx_path or not os.path.exists(docx_path):
        return ""
    
    try:
        from docx import Document
        doc = Document(docx_path)
        
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Tablolardan da metin çıkar
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        
        return '\n'.join(paragraphs)
    except ImportError:
        print("⚠️ python-docx yüklü değil. DOCX desteği devre dışı.")
        return ""
    except Exception as e:
        print(f"❌ DOCX Hata: {docx_path} → {e}")
        return ""


def extract_text_with_ocr(pdf_path):
    """
    Görsel (taranmış) PDF'den OCR ile metin çıkarır.
    """
    try:
        from PIL import Image
        import pytesseract
        
        doc = fitz.open(pdf_path)
        full_text = ""
        
        for page_num in range(min(doc.page_count, 5)):  # Max 5 sayfa (performans)
            page = doc[page_num]
            # Sayfayı yüksek çözünürlüklü görüntüye çevir
            mat = fitz.Matrix(2, 2)  # 2x zoom
            pix = page.get_pixmap(matrix=mat)
            
            # PIL Image'e dönüştür
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            
            # OCR uygula (Türkçe + İngilizce)
            try:
                page_text = pytesseract.image_to_string(img, lang='tur+eng')
            except Exception:
                # Türkçe dil paketi yoksa sadece İngilizce dene
                try:
                    page_text = pytesseract.image_to_string(img, lang='eng')
                except Exception:
                    page_text = pytesseract.image_to_string(img)
            
            if page_text:
                full_text += page_text + "\n"
        
        doc.close()
        return full_text
        
    except ImportError:
        print("⚠️ pytesseract veya Pillow yüklü değil. OCR desteği devre dışı.")
        return ""
    except Exception as e:
        print(f"⚠️ OCR Hatası: {e}")
        return ""


def extract_text_auto(file_path):
    """
    Dosya uzantısına göre otomatik metin çıkarıcı.
    PDF, DOCX ve görsel PDF (OCR) destekler.
    """
    if not file_path or not os.path.exists(file_path):
        return ""
    
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    else:
        print(f"⚠️ Desteklenmeyen format: {ext}")
        return ""
